"""
Document Processor for UzSWLU Chatbot
PDF, Word, TXT fayllarni o'qib, ChromaDB ga qo'shadi.
Improved chunking for better RAG retrieval.
"""
import os
import re
import json
from typing import List, Dict, Optional
from datetime import datetime

# PDF processing
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Word processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

print(f"📄 Document Processor: PDF={PDF_AVAILABLE}, DOCX={DOCX_AVAILABLE}")


class DocumentProcessor:
    """Hujjatlarni qayta ishlash va chunk'larga bo'lish."""
    
    def __init__(self, chunk_size: int = 1200, chunk_overlap: int = 400):
        # Optimized for Qwen 3B: approx 400-500 tokens
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def process_file(self, file_path: str, doc_type: str = None, doc_title: str = "Knowledge Base") -> Dict:
        """Faylni o'qib, matn va metadatani qaytaradi. Yaxshilangan error handling."""
        if not os.path.exists(file_path):
            return {
                'text': '', 'chunks': [], 'metadata': {},
                'success': False, 'error': f'Fayl topilmadi: {file_path}'
            }
        
        if doc_type is None:
            doc_type = self._detect_file_type(file_path)
        
        try:
            print(f"📖 Fayl o'qilmoqda: {file_path} (turi: {doc_type})")
            
            if doc_type == 'pdf':
                if not PDF_AVAILABLE:
                    return {
                        'text': '', 'chunks': [], 'metadata': {},
                        'success': False, 'error': 'PyPDF2 kutubxonasi o\'rnatilmagan'
                    }
                text = self._read_pdf(file_path)
            elif doc_type == 'word':
                if not DOCX_AVAILABLE:
                    return {
                        'text': '', 'chunks': [], 'metadata': {},
                        'success': False, 'error': 'python-docx kutubxonasi o\'rnatilmagan'
                    }
                text = self._read_word(file_path)
            elif doc_type == 'text':
                text = self._read_text(file_path)
            else:
                return {
                    'text': '', 'chunks': [], 'metadata': {},
                    'success': False, 'error': f'Noma\'lum fayl turi: {doc_type}'
                }
            
            if not text or len(text.strip()) < 10:
                return {
                    'text': '', 'chunks': [], 'metadata': {},
                    'success': False, 'error': 'Fayl bo\'sh yoki matn topilmadi'
                }
            
            print(f"✅ Fayl o'qildi: {len(text)} belgi")
            
            # Matnni tozalash
            text = self._clean_text(text)
            print(f"🧹 Matn tozalandi: {len(text)} belgi")
            
            # Smart chunking
            print("🔪 Chunk'larga ajratilmoqda...")
            chunks = self._smart_chunking(text, doc_title)
            print(f"✅ {len(chunks)} ta chunk yaratildi")
            
            if not chunks:
                return {
                    'text': '', 'chunks': [], 'metadata': {},
                    'success': False, 'error': 'Chunk yaratilmadi - fayl juda qisqa yoki format noto\'g\'ri'
                }
            
            metadata = {
                'file_path': file_path,
                'file_name': os.path.basename(file_path),
                'doc_type': doc_type,
                'char_count': len(text),
                'chunk_count': len(chunks),
                'processed_at': datetime.now().isoformat()
            }
            
            return {
                'text': text, 'chunks': chunks, 'metadata': metadata,
                'success': True, 'error': None
            }
            
        except Exception as e:
            import traceback
            error_trace = traceback.format_exc()
            print(f"❌ Fayl qayta ishlashda xatolik: {e}")
            print(f"Traceback: {error_trace}")
            return {
                'text': '', 'chunks': [], 'metadata': {},
                'success': False, 'error': f'Xatolik: {str(e)}'
            }
    
    def _detect_file_type(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()
        type_map = {
            '.pdf': 'pdf', '.docx': 'word', '.doc': 'word',
            '.txt': 'text', '.md': 'text',
        }
        return type_map.get(ext, 'text')
    
    def _read_pdf(self, file_path: str) -> str:
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 kutubxonasi o'rnatilmagan")
        
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return '\n\n'.join(text_parts)
    
    def _read_word(self, file_path: str) -> str:
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx kutubxonasi o'rnatilmagan")
        
        doc = DocxDocument(file_path)
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)
        
        return '\n\n'.join(text_parts)
    
    def _read_text(self, file_path: str) -> str:
        encodings = ['utf-8', 'utf-16', 'cp1251', 'latin-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read()
            except UnicodeDecodeError:
                continue
        raise ValueError("Fayl kodlashini aniqlab bo'lmadi")
    
    def _clean_text(self, text: str) -> str:
        """Matnni tozalash - URL'lar va ortiqcha narsalarni olib tashlash."""
        """Matnni tozalash va jadvallarni saqlashga harakat qilish."""
        # Non-breaking space'larni oddiy bo'shliqqa almashtirish
        text = text.replace('\xa0', ' ')
        
        # Ketma-ket bo'shliqlarni (jadvallarni) pipe bilan almashtirish
        # 3 ta yoki undan ko'p bo'shliq jadval ustunlari oralig'i bo'lishi mumkin
        text = re.sub(r' {3,}', ' | ', text)
        
        # Boshqaruv belgilarini olib tashlash
        text = "".join(ch for ch in text if ch.isprintable() or ch in '\n\r\t')
        
        # Ketma-ket kelgan pipe'larni bittalash
        text = re.sub(r'(\| ){2,}', '| ', text)
        
        # University headers often have specific patterns
        # e.g., "O'ZBEKISTON RESPUBLIKASI OLIY TA'LIM..."
        text = re.sub(r'O\'ZBEKISTON RESPUBLIKASI[^.\n]*', '', text, flags=re.IGNORECASE)
        
        # [Apply Now] havolalarini olib tashlash
        text = re.sub(r'\[Apply[^\]]*\]\s*\([^)]+\)', '', text)
        
        # Ketma-ket kelgan bo'shliqlarni bitta bo'shliqqa, lekin yangi qatorlarni saqlash
        # \s+ barcha whitespace'larni (shu jumladan \n) bittaga aylantirib yuboradi
        # Shuning uchun faqat oddiy bo'shliqlarni bittalaymiz
        text = re.sub(r' {2,}', ' ', text)
        
        # Ortiqcha yangi qatorlarni tozalash (3+ tasi 2 ta qoladi)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Maxsus belgilarni tozalash 
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        return text.strip()
    
    def _extract_table_structure(self, text: str) -> dict:
        """
        v6.0: Extract table structure from text.
        Detects tables by looking for pipe separators and multiple columns.
        """
        lines = text.split('\n')
        table_data = {'headers': [], 'rows': []}
        
        # Look for lines with pipe separators (from our table preservation)
        table_lines = [line for line in lines if '|' in line and line.count('|') >= 2]
        
        if not table_lines:
            return table_data
        
        # First line with most pipes is likely the header
        if table_lines:
            header_line = max(table_lines, key=lambda x: x.count('|'))
            headers = [h.strip() for h in header_line.split('|') if h.strip()]
            table_data['headers'] = headers[:5]  # Limit to 5 columns
            
            # Extract row data (simple heuristic)
            for line in table_lines[1:10]:  # Limit to 10 rows
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if len(cells) >= 2:
                    row_dict = {}
                    for i, cell in enumerate(cells[:len(headers)]):
                        if i < len(headers):
                            row_dict[headers[i]] = cell
                    table_data['rows'].append(row_dict)
        
        return table_data
    
    def _smart_chunking(self, text: str, doc_title: str = "Knowledge Base") -> List[str]:
        """Smart chunking - ma'noli bo'laklarga ajratish. Yaxshilangan versiya."""
        if not text:
            return []
        
        # Matnni tozalash va paragrafga ajratish
        paragraphs = self._split_into_paragraphs(text)
        
        # Semantik bo'limlarni aniqlash (yaxshilangan)
        sections = self._identify_sections_improved(text, paragraphs)
        
        # Har bir bo'limni chunk'larga ajratish
        chunks = []
        for section_title, section_text in sections:
            section_chunks = self._chunk_section_improved(section_title, section_text, doc_title)
            chunks.extend(section_chunks)  # Now returns list of dicts
        
        # Bo'sh yoki juda qisqa chunk'larni olib tashlash
        chunks = [chunk for chunk in chunks if len(chunk['text'].strip()) >= 50]
        
        # Takrorlanishlarni olib tashlash
        # Note: _remove_duplicates will need adjustment if it expects strings
        chunks = self._remove_duplicates(chunks)
        
        return chunks
    
    def _split_into_paragraphs(self, text: str) -> List[str]:
        """Matnni paragraflarga ajratish."""
        # Paragraf ajratuvchilar
        paragraph_separators = [
            r'\n\s*\n',  # Ikki yoki ko'p bo'sh qator
            r'\n(?=[A-ZА-ЯЁ])',  # Katta harf bilan boshlanadigan qator
        ]
        
        paragraphs = []
        current_paragraph = ""
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                if current_paragraph:
                    paragraphs.append(current_paragraph)
                    current_paragraph = ""
            else:
                if current_paragraph:
                    current_paragraph += " " + line
                else:
                    current_paragraph = line
        
        if current_paragraph:
            paragraphs.append(current_paragraph)
        
        return paragraphs
    
    def _identify_sections_improved(self, text: str, paragraphs: List[str]) -> List[tuple]:
        """Yaxshilangan: Matnni semantik bo'limlarga ajratish - universal."""
        # Kengaytirilgan section markers (o'zbek, ingliz, rus tillari)
        markers = [
            # Umumiy bo'limlar
            (r'(?:^|\n)\s*(?:About|Haqida|О нас|About Us|Haqimizda)', 'About'),
            (r'(?:^|\n)\s*(?:Introduction|Kirish|Введение|Tanishuv)', 'Introduction'),
            (r'(?:^|\n)\s*(?:Overview|Umumiy|Обзор)', 'Overview'),
            
            # Ta'lim bo'limlari
            (r'(?:^|\n)\s*(?:Programs?|Dasturlar|Программы|Ta\'lim dasturlari)', 'Programs'),
            (r'(?:^|\n)\s*(?:Bachelor|Bakalavr|Бакалавр|Bakalavriat)', 'Bachelor Programs'),
            (r'(?:^|\n)\s*(?:Master|Magistr|Магистр|Magistratura)', 'Master Programs'),
            (r'(?:^|\n)\s*(?:PhD|Doktorantura|Докторантура)', 'PhD Programs'),
            (r'(?:^|\n)\s*(?:Faculties?|Fakultetlar|Факультеты)', 'Faculties'),
            (r'(?:^|\n)\s*(?:Departments?|Kafedralar|Кафедры)', 'Departments'),
            
            # Qabul bo'limlari
            (r'(?:^|\n)\s*(?:Admission|Qabul|Поступление|Qabul jarayoni)', 'Admission'),
            (r'(?:^|\n)\s*(?:Requirements?|Talablar|Требования)', 'Requirements'),
            (r'(?:^|\n)\s*(?:Documents?|Hujjatlar|Документы)', 'Documents'),
            (r'(?:^|\n)\s*(?:Deadline|Muddat|Срок)', 'Deadline'),
            (r'(?:^|\n)\s*(?:Application|Ariza|Заявка)', 'Application'),
            
            # To'lov bo'limlari
            (r'(?:^|\n)\s*(?:Tuition|Kontrakt|Стоимость|To\'lov)', 'Tuition'),
            (r'(?:^|\n)\s*(?:Fees?|To\'lovlar|Платежи)', 'Fees'),
            (r'(?:^|\n)\s*(?:Scholarship|Grant|Stipendiya|Стипендия)', 'Scholarship'),
            
            # Hamkorlik
            (r'(?:^|\n)\s*(?:Partners?|Hamkorlar|Партнеры)', 'Partners'),
            (r'(?:^|\n)\s*(?:International|Xalqaro|Международный)', 'International'),
            (r'(?:^|\n)\s*(?:Exchange|Almashinuv|Обмен)', 'Exchange'),
            
            # Boshqa
            (r'(?:^|\n)\s*(?:Contact|Aloqa|Контакты)', 'Contact'),
            (r'(?:^|\n)\s*(?:Location|Manzil|Адрес)', 'Location'),
            (r'(?:^|\n)\s*(?:Benefits?|Afzalliklar|Преимущества)', 'Benefits'),
            (r'(?:^|\n)\s*(?:Opportunities?|Imkoniyatlar|Возможности)', 'Opportunities'),
        ]
        
        sections = []
        used_ranges = []
        
        # Markerlarni qidirish
        for pattern, title in markers:
            matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
            for match in matches:
                start = match.start()
                
                # Keyingi marker yoki matn oxirigacha
                end = len(text)
                for next_pattern, _ in markers:
                    if next_pattern != pattern:
                        next_match = re.search(next_pattern, text[start+20:], re.IGNORECASE | re.MULTILINE)
                        if next_match:
                            potential_end = start + 20 + next_match.start()
                            if potential_end < end:
                                end = potential_end
                
                # Overlap tekshirish
                overlaps = False
                for used_start, used_end in used_ranges:
                    if start < used_end and end > used_start:
                        overlaps = True
                        break
                
                if not overlaps:
                    section_text = text[start:end].strip()
                    if len(section_text) > 50:  # Minimal uzunlik
                        sections.append((title, section_text))
                        used_ranges.append((start, end))
        
        # Paragraflar bo'yicha bo'limlar (marker topilmasa)
        if not sections:
            # Paragraflarni guruhlash
            current_section = []
            current_title = "Section 1"
            section_num = 1
            
            for para in paragraphs:
                if len(para) > 100:  # Minimal paragraf uzunligi
                    current_section.append(para)
                    
                    # Har 3-4 paragrafdan keyin yangi bo'lim
                    if len(current_section) >= 3:
                        section_text = "\n\n".join(current_section)
                        if len(section_text) > 200:
                            sections.append((current_title, section_text))
                            section_num += 1
                            current_title = f"Section {section_num}"
                            current_section = []
            
            # Qolgan paragraflar
            if current_section:
                section_text = "\n\n".join(current_section)
                if len(section_text) > 200:
                    sections.append((current_title, section_text))
        
        # Agar hali ham bo'lim topilmasa, butun matnni bo'lim sifatida qaytarish
        if not sections:
            # Matnni teng qismlarga bo'lish
            chunk_size = len(text) // 3
            for i in range(0, len(text), chunk_size):
                chunk = text[i:i+chunk_size].strip()
                if len(chunk) > 200:
                    sections.append((f"Part {i//chunk_size + 1}", chunk))
        
        # Pozitsiya bo'yicha tartiblash
        sections.sort(key=lambda x: text.find(x[1][:100]) if x[1] else 0)
        
        return sections

    def _recursive_split(self, text: str, delimiters: List[str], chunk_size: int, chunk_overlap: int) -> List[str]:
        """AnythingLLM uslubida: Matnni rekursiv ravishda bo'lish."""
        if len(text) <= chunk_size:
            return [text]
        
        # Agat delimiterlar qolmagan bo'lsa, shunchaki kesish
        if not delimiters:
            return [text[i:i + chunk_size] for i in range(0, len(text), chunk_size - chunk_overlap)]
        
        # Hozirgi delimiter bilan bo'lish
        delimiter = delimiters[0]
        new_delimiters = delimiters[1:]
        
        splits = text.split(delimiter)
        final_chunks = []
        current_chunk = ""
        
        for split in splits:
            if not split: continue
            
            # Agar bitta split o'zi chunk_size dan katta bo'lsa, uni keyingi delimiterga yuborish
            if len(split) > chunk_size:
                if current_chunk:
                    final_chunks.append(current_chunk.strip())
                    current_chunk = ""
                
                # Rekursiv bo'lish
                sub_chunks = self._recursive_split(split, new_delimiters, chunk_size, chunk_overlap)
                final_chunks.extend(sub_chunks)
                continue
            
            # Agar qo'shish mumkin bo'lsa
            if len(current_chunk) + len(split) + len(delimiter) <= chunk_size:
                current_chunk += (delimiter if current_chunk else "") + split
            else:
                # To'ldi, saqlaymiz
                if current_chunk:
                    final_chunks.append(current_chunk.strip())
                current_chunk = split
        
        if current_chunk:
            final_chunks.append(current_chunk.strip())
            
        return final_chunks
    
    def _chunk_section_improved(self, title: str, text: str, doc_title: str = "") -> List[dict]:
        """
        v6.0: Enhanced with table structure extraction.
        Returns list of dicts with 'text' and 'metadata' keys.
        """
        # Eng yaxshi delimiterlar iyerarxiyasi
        delimiters = ["\n\n", "\n", ". ", " ", ""]
        
        # Rekursiv bo'lish
        raw_chunks = self._recursive_split(text, delimiters, self.chunk_size, self.chunk_overlap)
        
        # v6.0: Extract table structure
        table_data = self._extract_table_structure(text)
        
        # Meta ma'lumotlar bilan boyitish
        final_chunks = []
        for i, chunk in enumerate(raw_chunks):
            if len(chunk.strip()) < 100: continue
            
            # Metadata labeling
            meta_header = f"### [HUJJAT: {doc_title} | BO'LIM: {title}]\n"
            
            # v6.0: Build metadata
            metadata = {
                'section_title': title,
                'parent_context': doc_title,
                'table_headers': table_data.get('headers', []),
                'row_data': table_data.get('rows', [])[i] if i < len(table_data.get('rows', [])) else {}
            }
            
            final_chunks.append({
                'text': f"{meta_header}{chunk.strip()}",
                'metadata': metadata
            })
            
        return final_chunks
    
    def _remove_duplicates(self, chunks: List[dict]) -> List[dict]:
        """Takrorlanuvchi chunk'larni olib tashlash."""
        unique_chunks = []
        seen_content = []
        
        for chunk_dict in chunks:
            chunk = chunk_dict['text']
            is_duplicate = False
            chunk_words = set(chunk.lower().split())
            
            for seen_words in seen_content:
                # Calculate similarity
                if chunk_words and seen_words:
                    intersection = len(chunk_words & seen_words)
                    union = len(chunk_words | seen_words)
                    similarity = intersection / union if union > 0 else 0
                    
                    if similarity > 0.7:  # 70% o'xshash - duplicate
                        is_duplicate = True
                        break
            
            if not is_duplicate:
                unique_chunks.append(chunk_dict)
                seen_content.append(chunk_words)
        
        return unique_chunks


class DocumentRAGIntegration:
    """Hujjatlarni RAG tizimiga ulash."""
    
    def __init__(self):
        # Chunk size: 300-500 tokens (approx 1500 chars for Uzbek)
        # Chunk overlap: 50-100 tokens (approx 300 chars)
        self.processor = DocumentProcessor(chunk_size=1500, chunk_overlap=300)
    
    def process_and_store(self, document) -> Dict:
        """Document modelini qayta ishlab, ChromaDB ga saqlash. Yaxshilangan versiya."""
        from rag_service import RAGService
        from django.utils import timezone
        import os
        from langdetect import detect
        
        try:
            if document.file_path:
                file_path = document.file_path.path
                # Map source_type to what processor expects
                doc_type_map = {'pdf': 'pdf', 'doc': 'word', 'text': 'text', 'html': 'text'}
                doc_type = doc_type_map.get(document.source_type, 'text')
                
                # File path'ni tekshirish
                if not os.path.exists(file_path):
                    error_msg = f"Fayl topilmadi: {file_path}"
                    print(f"❌ {error_msg}")
                    return {
                        'success': False,
                        'chunks_created': 0,
                        'error': error_msg
                    }
            else:
                return {
                    'success': False,
                    'chunks_created': 0,
                    'error': 'Fayl yo\'li topilmadi'
                }
            
            print(f"📄 Hujjat qayta ishlanmoqda: {document.title} ({doc_type})")
            document.status = 'processing'
            document.save(update_fields=['status'])
            
            # Process file
            result = self.processor.process_file(file_path, doc_type, document.title)
            
            if not result['success']:
                print(f"❌ Fayl qayta ishlashda xatolik: {result['error']}")
                return {
                    'success': False,
                    'chunks_created': 0,
                    'error': result['error']
                }
            
            chunks = result['chunks']
            if not chunks:
                print(f"⚠️ Chunk yaratilmadi: {document.title}")
                return {
                    'success': False,
                    'chunks_created': 0,
                    'error': 'Chunk yaratilmadi - fayl bo\'sh yoki qayta ishlab bo\'lmaydi'
                }
            
            # Detect language if not explicitly set to something other than default 'uz'
            full_text = result['text']
            try:
                detected_lang = detect(full_text[:2000])
                if detected_lang in ['uz', 'ru', 'en']:
                    document.lang = detected_lang
                    print(f"🌐 Aniqlangan til: {detected_lang}")
            except:
                print("⚠️ Tilni aniqlab bo'lmadi, default 'uz' qoldiriladi")
            
            print(f"✅ {len(chunks)} ta chunk yaratildi")
            
            # Store in ChromaDB
            print(f"💾 ChromaDB ga saqlanmoqda...")
            rag = RAGService()
            
            # Delete old chunks for this document from Chroma
            try:
                existing = rag.collection.get(
                    where={"document_id": str(document.id)}
                )
                if existing and existing['ids']:
                    rag.collection.delete(ids=existing['ids'])
                    print(f"🗑️ {len(existing['ids'])} ta eski chunk ChromaDB'dan o'chirildi")
            except Exception as e:
                print(f"⚠️ Eski chunk'larni o'chirishda xatolik: {e}")
            
            # Add new chunks to Chroma
            ids = [f"doc_{document.id}_chunk_{i}" for i in range(len(chunks))]
            # Infer category from title
            category = 'Umumiy'
            title_lower = document.title.lower()
            if any(k in title_lower for k in ['kontrakt', 'to\'lov', 'shartnoma']): category = 'Kontrakt'
            elif any(k in title_lower for k in ['yotoqxona', 'ttj', 'turar joy']): category = 'Talaba hayoti'
            elif any(k in title_lower for k in ['qabul', 'imtihon', 'imtixon']): category = 'Qabul'
            elif any(k in title_lower for k in ['magistr']): category = 'Magistratura'

            # Extract text and metadata for ChromaDB
            document_texts = [c['text'] for c in chunks]
            
            metadatas = []
            for i, c in enumerate(chunks):
                meta = {
                    'document_id': str(document.id),
                    'title': document.title,
                    'chunk_index': i,
                    'source': 'document',
                    'lang': document.lang,
                    'type': 'document',
                    'category': category
                }
                # Merge with processor metadata and sanitize for ChromaDB (no lists/dicts)
                proc_meta = c.get('metadata', {})
                for k, v in proc_meta.items():
                    if isinstance(v, (list, dict)):
                        meta[k] = json.dumps(v)
                    else:
                        meta[k] = v
                metadatas.append(meta)
            
            # Add new chunks to Chroma
            # v6.1: Manual embedding for nomic prefix support
            if hasattr(rag.embedding_fn, 'model_name') and "nomic" in rag.embedding_fn.model_name:
                embeddings = rag.embedding_fn(document_texts, prefix="search_document: ")
                rag.collection.add(
                    documents=document_texts,
                    embeddings=embeddings,
                    metadatas=metadatas,
                    ids=ids
                )
            else:
                rag.collection.add(
                    documents=document_texts,
                    metadatas=metadatas,
                    ids=ids
                )
            print(f"✅ {len(chunks)} ta chunk ChromaDB ga qo'shildi")
            
            # Store metadata and text in PostgreSQL
            from chatbot_app.models import DocumentChunk
            DocumentChunk.objects.filter(document=document).delete()
            
            chunk_objs = [
                DocumentChunk(
                    document=document,
                    chunk_index=i,
                    embedding_id=ids[i],
                    chunk_text=chunks[i]['text'],
                    lang=document.lang,
                    metadata={**metadatas[i], 'title': document.title, 'index': i}
                ) for i in range(len(chunks))
            ]
            DocumentChunk.objects.bulk_create(chunk_objs)
            print(f"✅ {len(chunks)} ta chunk metadata PostgreSQL ga saqlandi")
            
            # Update Document status
            document.status = 'ready'
            document.save(update_fields=['status', 'lang'])
            
            return {
                'success': True,
                'chunks_created': len(chunks),
                'error': None
            }
            
        except Exception as e:
            import traceback
            print(f"❌ Hujjat qayta ishlashda xatolik: {e}\n{traceback.format_exc()}")
            
            document.status = 'failed'
            document.save(update_fields=['status'])
            
            return {
                'success': False,
                'chunks_created': 0,
                'error': str(e)
            }


# Utility functions
def process_document(document_id: int) -> Dict:
    """Helper function to process a document by ID."""
    try:
        from chatbot_app.models import Document
        document = Document.objects.get(id=document_id)
        integration = DocumentRAGIntegration()
        return integration.process_and_store(document)
    except Exception as e:
        return {'success': False, 'chunks_created': 0, 'error': str(e)}


def reprocess_all_documents() -> Dict:
    """Reprocess all documents."""
    try:
        from chatbot_app.models import Document
        # Use status instead of is_active
        documents = Document.objects.filter(status='ready')
        
        results = {
            'total': documents.count(),
            'success': 0,
            'failed': 0,
            'errors': []
        }
        
        for doc in documents:
            result = process_document(doc.id)
            if result['success']:
                results['success'] += 1
            else:
                results['failed'] += 1
                results['errors'].append(f"{doc.title}: {result['error']}")
        
        return results
    except Exception as e:
        return {'total': 0, 'success': 0, 'failed': 0, 'errors': [str(e)]}


# Global instance for easy import
doc_rag = DocumentRAGIntegration()
